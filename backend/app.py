import os
import google.generativeai as genai
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document  # Keep this for Word export
import io
from dotenv import load_dotenv
import unicodedata  # Good practice for text cleaning, though less critical for docx
import re  # Good practice for text cleaning

load_dotenv()

app = Flask(__name__)
CORS(app)

# Configure Gemini API
try:
    gemini_api_key = os.getenv("GOOGLE_API_KEY")
    if not gemini_api_key:
        raise ValueError("GOOGLE_API_KEY not found in environment variables.")
    genai.configure(api_key=gemini_api_key)
    model = genai.GenerativeModel("gemini-2.0-flash")  # CONFIRMED MODEL NAME
except Exception as e:
    print(f"Error configuring Gemini API: {e}")
    model = None


def generate_mcqs_with_gemini(topic_name, chapter_content, pyqs):
    if not model:
        return "Error: Gemini model not initialized. Please check API key and configuration."

    prompt = f"""
    You are an expert MCQ generator for NCERT Social Science topics.
    Generate 5 high-quality Multiple Choice Questions (MCQs) based on the following topic and content.
    If no content is provided, generate general MCQs for the topic.

    Topic: "{topic_name}"

    Chapter Content (use this as the primary source):
    ---
    {chapter_content if chapter_content else "No specific content provided. Generate general MCQs for the topic."}
    ---

    Previous Year Questions (PYQs) for reference (optional, use for style/focus if relevant):
    ---
    {pyqs if pyqs else "No PYQs provided."}
    ---

    Each MCQ should strictly follow this format:
    Q[Number]. [Question Text]?
    A) [Option A]
    B) [Option B]
    C) [Option C]
    D) [Option D]
    Answer: [Correct Option Letter, e.g., B]

    Ensure each MCQ is distinct, clear, and directly related to the topic/content.
    Provide only the MCQs and their answers, with no extra introductory or concluding text.
    Start directly with "Q1."
    """

    try:
        response = model.generate_content(prompt)
        if response.parts:
            return response.text
        else:
            if response.candidates and response.candidates[0].content.parts:
                return response.candidates[0].content.parts[0].text
            return "Error: No content generated or unexpected response structure from Gemini."
    except Exception as e:
        print(f"Error during Gemini API call: {e}")
        if hasattr(e, "response") and hasattr(e.response, "prompt_feedback"):
            return f"Error generating MCQs: Content Filtered by API. Prompt Feedback: {e.response.prompt_feedback}"
        return f"Error generating MCQs: {str(e)}"


# Helper function to sanitize text for filenames
def sanitize_filename(name):
    # Replace spaces with underscores, and strip invalid characters
    clean_name = name.replace(" ", "_").replace("/", "-").replace("\\", "")
    clean_name = "".join(x for x in clean_name if x.isalnum() or x in "._-").strip()
    clean_name = clean_name.replace("__", "_").replace("--", "-")
    return clean_name


@app.route("/api/generate-mcqs", methods=["POST"])
def api_generate_mcqs():
    data = request.json
    topic_name = data.get("topicName", "")
    chapter_content = data.get("chapterContent", "")
    pyqs = data.get("pyqs", "")

    if not topic_name:
        return jsonify({"error": "Topic name is required"}), 400
    if not model:
        return (
            jsonify({"error": "Gemini model not initialized. Check backend logs."}),
            500,
        )

    mcqs_text = generate_mcqs_with_gemini(topic_name, chapter_content, pyqs)

    if "Error:" in mcqs_text:
        return jsonify({"error": mcqs_text}), 500

    return jsonify({"mcqs": mcqs_text})


@app.route("/api/export-word", methods=["POST"])
def export_word():
    data = request.json
    mcqs_text = data.get("mcqs", "")
    topic_name_for_export = data.get("topicName", "").strip()

    document = Document()
    title = (
        f"{topic_name_for_export} MCQs" if topic_name_for_export else "Generated MCQs"
    )
    document.add_heading(title, 0)

    lines = mcqs_text.split("\n")
    for line in lines:
        if line.strip():
            document.add_paragraph(line)

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    download_filename = (
        f"{sanitize_filename(topic_name_for_export)}_mcqs.docx"
        if topic_name_for_export
        else "mcqs.docx"
    )

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=download_filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# Removed the /api/export-pdf route

if __name__ == "__main__":
    app.run(debug=True)
